# HW2_app.py
# NYT Articles analysis: understand mass shooting coverage by state.
# Layout variant of HW1_app.py (tabs + header + global CSS); server logic unchanged.
# === TEMPORARY DEPLOYMENT DIAGNOSIS - REMOVE AFTER ===
import os
from pathlib import Path as _Path

print("=== DEPLOYMENT DIAGNOSIS ===")
print(f"Working directory: {os.getcwd()}")
print(f"__file__ location: {_Path(__file__).resolve()}")
print(f"Files in cwd: {sorted(os.listdir('.'))}")

_repo_root = _Path(__file__).resolve().parent.parent
print(f"Repo root: {_repo_root}")
print(f"Contents of repo root: {sorted(os.listdir(_repo_root))}")

_cache = _repo_root / "HOMEWORK_1" / "nyt_2025_shootings_cache.json"
_gva = _repo_root / "HOMEWORK_1" / "gva_data.csv"

print(f"Cache path tried: {_cache}")
print(f"Cache exists: {_cache.exists()}")
if _cache.exists():
    print(f"Cache size (bytes): {_cache.stat().st_size}")
    import json as _json
    with open(_cache) as _f:
        _cache_data = _json.load(_f)
    print(f"Cache article count: {len(_cache_data)}")
    _nv = [
        a for a in _cache_data
        if "nevada" in (
            a.get("abstract", "") +
            a.get("headline", {}).get("main", "")
        ).lower()
    ]
    print(f"Nevada articles in cache: {len(_nv)}")
else:
    print("CACHE MISSING - this is the problem")

print(f"GVA path tried: {_gva}")
print(f"GVA exists: {_gva.exists()}")
if _gva.exists():
    print(f"GVA size (bytes): {_gva.stat().st_size}")
    import pandas as _pd
    _gva_df = _pd.read_csv(_gva)
    print(f"GVA total rows: {len(_gva_df)}")
    print(f"GVA 2025 rows: {len(_gva_df[_gva_df['year']==2025])}")
    _nv_gva = _gva_df[
        (_gva_df['year']==2025) &
        (_gva_df['state']=='Nevada')
    ]
    print(f"Nevada 2025 GVA events: {len(_nv_gva)}")
    if len(_nv_gva) > 0:
        print(_nv_gva[['city_or_county','date_fixed']].to_string())
else:
    print("GVA MISSING - this is the problem")

print("=== END DIAGNOSIS ===")
# === END TEMPORARY DIAGNOSIS ===
from collections import defaultdict
import base64
import json
from pathlib import Path
import sys
import traceback
from typing import Any, Dict, List, Optional

# HOMEWORK_2 on sys.path first so we can import homework1_dir (must match GVA/NYT paths in HW2_multi_agent).
_hw2 = Path(__file__).resolve().parent
if str(_hw2) not in sys.path:
    sys.path.insert(0, str(_hw2))

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from shiny import App, Inputs, Outputs, Session, reactive, render, ui
from shiny.ui import Progress

from HW2_multi_agent import (
    HIGH_PROFILE_ARTICLE_THRESHOLD,
    STATE_NAME_TO_ABBR,
    build_gva_events_for_states,
    homework1_dir,
    run_agent_pipeline,
)

_HOMEWORK_1 = Path(homework1_dir())
_HOMEWORK_2 = Path(__file__).resolve().parent
for p in (_HOMEWORK_1, _HOMEWORK_2):
    if str(p) not in sys.path:
        sys.path.insert(0, str(p))

_ALL_STATES_ABBR = list(STATE_NAME_TO_ABBR.values())
from HW1_state_analysis import (
    get_states_2025,
    run_state_analysis,
)
from HW1_nyt_cache import load_or_build_2025_cache


def compute_national_stats(events):
    """Per-state GVA totals, NYT-covered counts, coverage %, outlier counts (5+ matched articles)."""
    abbr_to_name = {v: k.title() for k, v in STATE_NAME_TO_ABBR.items()}
    stats = defaultdict(lambda: {"total": 0, "covered": 0, "outliers": 0})
    for event in events:
        s = str(event.get("state", "")).strip().upper()
        if not s:
            continue
        stats[s]["total"] += 1
        if event.get("matched_article_urls"):
            stats[s]["covered"] += 1
        if event.get("is_outlier"):
            stats[s]["outliers"] += 1
    for abbr in STATE_NAME_TO_ABBR.values():
        if abbr not in stats:
            stats[abbr] = {"total": 0, "covered": 0, "outliers": 0}
    rows = []
    for abbr, counts in stats.items():
        total = counts["total"]
        covered = counts["covered"]
        rows.append(
            {
                "state_abbr": abbr,
                "state_name": abbr_to_name.get(abbr, abbr),
                "total_events": total,
                "covered_events": covered,
                "coverage_rate": round(covered / total * 100, 1) if total > 0 else 0.0,
                "outlier_count": counts["outliers"],
            }
        )
    return pd.DataFrame(rows)


# National stats + hw2_events_all are computed in the server after NYT cache load (init_cache),
# so they use the same on-disk cache as load_or_build_2025_cache — not import-time (fixes Connect drift).

# --- HW2 ADDITION: pipeline startup ---
# RAG loaded once at startup across all states (reactive per-state use later).
rag_conn = None
try:
    from rag_setup import setup_rag

    rag_conn = setup_rag(
        cache_path=_HOMEWORK_1 / "nyt_2025_shootings_cache.json",
        states=list(STATE_NAME_TO_ABBR.values()),
        state_name_to_abbr=STATE_NAME_TO_ABBR,
        data_dir=_HOMEWORK_2 / "data",
        rebuild=False,
    )
    print("[Startup] RAG index loaded.", flush=True)
except Exception as _rag_err:
    print(
        f"[Startup] RAG index failed: {_rag_err} "
        f"(interpreter={sys.executable})",
        flush=True,
    )
# --- END HW2 ADDITION ---


def _state_pick_to_abbr(state_name: Optional[str]) -> Optional[str]:
    """Map select value (e.g. 'Ohio') to GVA/NYT event state code (e.g. 'OH')."""
    if not state_name:
        return None
    return STATE_NAME_TO_ABBR.get(str(state_name).strip().upper())


def _synthetic_validated_from_gva_events(events: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Agent 2 keys coverage off Agent-1-shaped rows. Derive them from GVA events that
    already have matched NYT URLs (same logic as the cache-backed comparison).
    """
    out: List[Dict[str, Any]] = []
    for ev in events:
        urls = ev.get("matched_article_urls") or []
        if not isinstance(urls, list) or not urls:
            continue
        u0 = str(urls[0]).strip()
        if not u0:
            continue
        out.append(
            {
                "relevant": True,
                "url": u0,
                "reason": "GVA/NYT cache match",
                "event": {
                    "city": ev.get("city"),
                    "state": str(ev.get("state", "")).strip().upper(),
                    "date": ev.get("date"),
                },
            }
        )
    return out


def _write_hw2_pipeline_docx(
    out_path: Path,
    title: str,
    body: str,
    stats_data: Optional[Dict[str, Any]] = None,
    bullet_points: Optional[List[str]] = None,
    outlier_events: Optional[List[Dict[str, Any]]] = None,
) -> None:
    """
    Write Agent 4 narrative plus structured statistical summary (real Word table),
    optional outliers and key findings. Omits internal prompt text (e.g. PRECOMPUTED / do not recalculate).
    """
    from docx import Document

    doc = Document()
    doc.add_heading(title, 0)
    for block in body.replace("\r\n", "\n").split("\n\n"):
        t = block.strip()
        if t:
            doc.add_paragraph(t)

    if not stats_data:
        doc.save(str(out_path))
        return

    sa = stats_data["state_a"]
    sb = stats_data["state_b"]
    sa_data = stats_data["state_a_data"]
    sb_data = stats_data["state_b_data"]
    demo_a = sa_data.get("demographics") if isinstance(sa_data.get("demographics"), dict) else {}
    demo_b = sb_data.get("demographics") if isinstance(sb_data.get("demographics"), dict) else {}

    def fmt_rate(v):
        if v is None:
            return "N/A"
        try:
            return f"{float(v):.1f}%"
        except (TypeError, ValueError):
            return "N/A"

    def fmt_currency(v):
        if v is None:
            return "N/A"
        try:
            return f"${float(v):,.0f}"
        except (TypeError, ValueError):
            return "N/A"

    def fmt_days(v):
        if v is None:
            return "N/A"
        try:
            return f"{float(v):.1f} days"
        except (TypeError, ValueError):
            return "N/A"

    def fmt_pop(v):
        if v is None:
            return "N/A"
        try:
            return f"{float(v):,.0f}"
        except (TypeError, ValueError):
            return "N/A"

    rows_data = [
        ("Total Shootings", str(sa_data.get("total_events", "")), str(sb_data.get("total_events", ""))),
        ("NYT Covered", str(sa_data.get("covered_events", "")), str(sb_data.get("covered_events", ""))),
        ("Coverage Rate", fmt_rate(sa_data.get("coverage_rate")), fmt_rate(sb_data.get("coverage_rate"))),
        (
            "Avg Days to First Article",
            fmt_days(sa_data.get("avg_days_to_first_article")),
            fmt_days(sb_data.get("avg_days_to_first_article")),
        ),
        (
            "Same-Day Coverage",
            fmt_rate(sa_data.get("same_day_coverage_pct")),
            fmt_rate(sb_data.get("same_day_coverage_pct")),
        ),
        (
            "Median Household Income",
            fmt_currency(demo_a.get("median_household_income")),
            fmt_currency(demo_b.get("median_household_income")),
        ),
        (
            "% White Population",
            fmt_rate(demo_a.get("pct_white")),
            fmt_rate(demo_b.get("pct_white")),
        ),
        (
            "State Population",
            fmt_pop(demo_a.get("population")),
            fmt_pop(demo_b.get("population")),
        ),
    ]

    doc.add_heading("Statistical summary", 1)
    nrows = 1 + len(rows_data)
    table = doc.add_table(rows=nrows, cols=3)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "Metric"
    h[1].text = str(sa)
    h[2].text = str(sb)
    for i, (metric, va, vb) in enumerate(rows_data, start=1):
        row = table.rows[i].cells
        row[0].text = metric
        row[1].text = va
        row[2].text = vb

    corr = stats_data.get("correlation_flag") or ""
    if corr:
        doc.add_paragraph(f"Correlation note: {corr}")
    speed = stats_data.get("speed_note") or ""
    if speed:
        doc.add_paragraph(f"Coverage speed: {speed}")

    if outlier_events:
        doc.add_heading("High-profile outlier events", 2)
        doc.add_paragraph(
            "These events had five or more matched NYT articles and may affect coverage rates."
        )
        for e in outlier_events:
            doc.add_paragraph(
                f"{e.get('city', '?')}, {e.get('state', '?')} ({e.get('date', '?')}) — "
                f"{e.get('article_count', 0)} NYT articles matched"
            )

    if bullet_points:
        doc.add_heading("Key findings", 2)
        for b in bullet_points:
            doc.add_paragraph(b, style="List Bullet")

    doc.save(str(out_path))


def _national_map_figure_bundle(
    df: pd.DataFrame,
    *,
    load_error: Optional[str] = None,
    is_loading: bool = False,
) -> dict:
    """
    Build the U.S. choropleth. Returns {"figure", "error", "diagnostics"}.
    On failure, figure is a placeholder with a short title; error holds the full traceback.
    """
    import plotly

    diag_lines = [
        f"plotly version: {getattr(plotly, '__version__', '?')}",
        f"DataFrame rows: {len(df)}",
        f"columns present: {list(df.columns)}",
    ]
    if load_error:
        diag_lines.append(f"national stats error: {load_error}")

    def _fallback_fig(title: str, subtitle: str = "") -> go.Figure:
        fig = go.Figure()
        title_text = title if not subtitle else f"{title} — {subtitle}"
        fig.update_layout(
            title={"text": title_text},
            height=460,
            paper_bgcolor="#f9fafb",
            margin={"r": 0, "t": 60, "l": 0, "b": 0},
        )
        return fig

    if is_loading:
        diag_lines.append("reason: waiting for NYT cache + GVA match load")
        return {
            "figure": _fallback_fig("National Picture", "Loading coverage data after NYT cache…"),
            "error": None,
            "diagnostics": "\n".join(diag_lines),
        }

    if df.empty or "state_abbr" not in df.columns:
        diag_lines.append("reason: empty DataFrame or missing state_abbr")
        return {
            "figure": _fallback_fig("National data unavailable", "No rows or missing columns."),
            "error": None,
            "diagnostics": "\n".join(diag_lines),
        }

    try:
        df = df.assign(
            hover_label=df.apply(
                lambda r: (
                    f"<b>{r['state_name']}</b><br>"
                    f"Total shootings: {int(r['total_events'])}<br>"
                    f"NYT covered: {int(r['covered_events'])}<br>"
                    f"Coverage rate: {r['coverage_rate']}%"
                    + (
                        f"<br>Outlier events (≥{HIGH_PROFILE_ARTICLE_THRESHOLD} articles): {int(r['outlier_count'])}"
                        if r["outlier_count"] > 0
                        else ""
                    )
                ),
                axis=1,
            )
        )

        fig = px.choropleth(
            df,
            locations="state_abbr",
            color="coverage_rate",
            locationmode="USA-states",
            scope="usa",
            color_continuous_scale=[
                (0.0, "#ffffff"),
                (0.05, "#dbeafe"),
                (0.35, "#60a5fa"),
                (0.7, "#2563eb"),
                (1.0, "#1e3a8a"),
            ],
            range_color=(0, max(100.0, float(df["coverage_rate"].max() or 1.0))),
            labels={"coverage_rate": "Coverage %"},
            height=460,
        )
        fig.update_traces(
            hovertemplate="%{hovertext}<extra></extra>",
            hovertext=df["hover_label"],
            marker_line_color="#ffffff",
            marker_line_width=0.8,
        )
        fig.update_geos(
            scope="usa",
            projection_type="albers usa",
            showlakes=True,
            lakecolor="#ffffff",
            bgcolor="#f9fafb",
            landcolor="#f1f5f9",
            subunitcolor="#e2e8f0",
            showsubunits=True,
        )
        fig.update_layout(
            margin={"r": 0, "t": 10, "l": 0, "b": 0},
            paper_bgcolor="#f9fafb",
            coloraxis_colorbar=dict(title="Coverage %"),
        )
        diag_lines.append(f"trace types: {[getattr(t, 'type', '?') for t in fig.data]}")
        diag_lines.append(f"coverage_rate min/max: {df['coverage_rate'].min()} / {df['coverage_rate'].max()}")
        return {"figure": fig, "error": None, "diagnostics": "\n".join(diag_lines)}
    except Exception:
        tb = traceback.format_exc()
        diag_lines.append("exception while building choropleth (see error block below)")
        return {
            "figure": _fallback_fig("Map build failed", "See error details below the plot area."),
            "error": tb,
            "diagnostics": "\n".join(diag_lines),
        }


def _plotly_figure_iframe_ui(fig: go.Figure, height_px: int = 460):
    """
    Show a Plotly figure in an iframe via a standalone HTML document (data: URL).
    shinywidgets' FigureWidget path often renders an empty plot in Shiny; this uses Plotly.js from CDN instead.
    """
    import plotly.io as pio

    fig.update_layout(height=height_px)
    html = pio.to_html(
        fig,
        full_html=True,
        include_plotlyjs="cdn",
        config={"responsive": True, "displayModeBar": True},
    )
    b64 = base64.b64encode(html.encode("utf-8")).decode("ascii")
    src = f"data:text/html;base64,{b64}"
    return ui.tags.iframe(
        src=src,
        title="U.S. mass shooting coverage by state",
        width="100%",
        height=f"{height_px}px",
        style="border: none; width: 100%; display: block; background: #f9fafb;",
    )


# Design tokens (UI only)
_HW_BG = "#f9fafb"
_HW_CARD = "#ffffff"
_HW_ACCENT = "#2563eb"
_HW_BORDER = "#e5e7eb"
_HW_TEXT = "#111827"
_HW_MUTED = "#6b7280"
_HW_RADIUS = "12px"
_HW_SHADOW = "0 4px 20px rgba(15, 23, 42, 0.06)"


def _takeaway_card(state_name: str, data: dict):
    """
    One takeaway card: Covered by NYT (event label + link + headline + pub date per event), Not covered (list).
    reported_events: list of {"label": "City, date", "articles": [{"url", "headline", "pub_date"}, ...]}.
    """
    reported = data.get("reported_events", [])
    not_reported = data.get("not_reported_events", [])
    max_shown = 15
    shown_reported = reported[:max_shown]
    shown_not = not_reported[:max_shown]

    def _one_event(entry):
        if isinstance(entry, str):
            return ui.div(ui.span(entry, class_="text-secondary"), class_="mb-2")
        label = entry.get("label", "")
        articles = entry.get("articles", [])
        kids = [ui.p(label, class_="fw-semibold mb-1", style="font-size: 0.95rem;")]
        article_lines = []
        for art in articles:
            url = art.get("url", "")
            headline = art.get("headline", "") or "(No headline)"
            pub_date = art.get("pub_date", "")
            line_parts = []
            if url:
                line_parts.append(ui.a(headline, href=url, target="_blank", rel="noopener"))
            else:
                line_parts.append(ui.span(headline, class_="text-muted"))
            if pub_date:
                line_parts.append(ui.span(f" (published {pub_date})", class_="text-muted small"))
            article_lines.append(ui.p(*line_parts, class_="small mb-0 ms-2", style="line-height: 1.4;"))
        # Scrollable list of all NYT articles for this shooting
        articles_container = ui.div(
            *article_lines,
            class_="border-start border-2 border-secondary ps-2 mb-3",
            style="max-height: 220px; overflow-y: auto; overflow-x: hidden;",
        )
        kids.append(articles_container)
        return ui.div(*kids)

    covered_list = [ui.div("None", class_="text-muted")] if not shown_reported else [_one_event(e) for e in shown_reported]
    if not shown_not:
        not_list = [ui.p("None", class_="text-muted mb-0")]
    else:
        not_list = [ui.tags.ul(*[ui.tags.li(ev, class_="small text-muted") for ev in shown_not], class_="mb-0 ps-3", style="line-height: 1.5;")]
    if len(not_reported) > max_shown:
        not_list.append(ui.p(f"... and {len(not_reported) - max_shown} more.", class_="small text-muted mt-1"))
    if len(reported) > max_shown:
        covered_list.append(ui.p(f"... and {len(reported) - max_shown} more events.", class_="small text-muted mt-1"))

    return ui.card(
        ui.card_header(
            ui.div(
                ui.h5(
                    f"{state_name}",
                    class_="mb-1",
                    style=f"color: {_HW_TEXT}; font-weight: 600; font-size: 1.05rem;",
                ),
                ui.p("Event-level detail", class_="mb-0 small", style=f"color: {_HW_MUTED};"),
            ),
            class_="border-0 pb-2",
            style=f"background: {_HW_CARD}; border-bottom: 1px solid {_HW_BORDER} !important;",
        ),
        ui.card_body(
            ui.h6(
                "Covered by NYT",
                class_="mb-2",
                style=f"font-size: 0.8rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.04em; color: {_HW_ACCENT};",
            ),
            ui.div(*covered_list),
            ui.h6(
                "Not covered",
                class_="mb-2 mt-3",
                style=f"font-size: 0.8rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.04em; color: {_HW_MUTED};",
            ),
            ui.div(*not_list),
            class_="pt-3",
        ),
        class_="col-12 hw-surface-card",
    )


def _comparison_table_rows(result: dict):
    """Build list of (metric, state_a_value, state_b_value). Coverage days as one heading row + min/max/mean/median sub-rows."""
    a = result.get("state_a") or {}
    b = result.get("state_b") or {}

    def _fmt(v):
        if v is None:
            return "—"
        if isinstance(v, float):
            return f"{v:.1f}"
        return str(v)

    return [
        ("Total mass shootings (2025)", _fmt(a.get("total_shootings")), _fmt(b.get("total_shootings"))),
        ("Events with ≥1 NYT article", _fmt(a.get("reported_count")), _fmt(b.get("reported_count"))),
        ("% reported by NYT", f"{_fmt(a.get('pct_reported'))}%", f"{_fmt(b.get('pct_reported'))}%"),
        ("Coverage duration (days)", "—", "—"),  # header row
        ("  min", _fmt(a.get("coverage_min_days")), _fmt(b.get("coverage_min_days"))),
        ("  max", _fmt(a.get("coverage_max_days")), _fmt(b.get("coverage_max_days"))),
        ("  mean", _fmt(a.get("coverage_mean_days")), _fmt(b.get("coverage_mean_days"))),
        ("  median", _fmt(a.get("coverage_median_days")), _fmt(b.get("coverage_median_days"))),
    ]


def _state_summary_card(state_name: str, a: dict):
    """Side-by-side summary: coverage % as focal metric + shootings + reported events."""
    total = a.get("total_shootings", 0)
    pct = a.get("pct_reported", 0)
    reported = a.get("reported_count", 0)
    return ui.div(
        ui.p(state_name, class_="hw-metric-state-name"),
        ui.div(
            ui.span(str(pct), class_="hw-metric-hero"),
            ui.span("%", class_="hw-metric-hero-suffix"),
            class_="d-flex align-items-baseline mb-2",
        ),
        ui.p("Share of GVA incidents with ≥1 NYT article", class_="hw-metric-caption mb-3"),
        ui.div(
            ui.div(
                ui.span("Total shootings", class_="hw-stat-label"),
                ui.span(str(total), class_="hw-stat-value"),
            ),
            ui.div(
                ui.span("Events w/ NYT", class_="hw-stat-label"),
                ui.span(str(reported), class_="hw-stat-value"),
            ),
            class_="d-flex justify-content-between gap-3 flex-wrap",
        ),
        class_="hw-metric-card",
    )


# 1. UI ############################


def _app_global_css() -> ui.Tag:
    return ui.tags.style(
        f"""
    :root {{
      --hw-bg: {_HW_BG};
      --hw-card: {_HW_CARD};
      --hw-accent: {_HW_ACCENT};
      --hw-border: {_HW_BORDER};
      --hw-text: {_HW_TEXT};
      --hw-muted: {_HW_MUTED};
      --hw-radius: {_HW_RADIUS};
      --hw-shadow: {_HW_SHADOW};
    }}
    body {{ background: var(--hw-bg) !important; color: var(--hw-text); }}
    .bslib-page-title {{ display: none; }}
    .hw-text-muted {{ color: var(--hw-muted) !important; }}
    .hw-hr {{ border-color: var(--hw-border); opacity: 1; }}
    .value-box {{
      border-radius: var(--hw-radius) !important;
      border: 1px solid var(--hw-border) !important;
      box-shadow: var(--hw-shadow) !important;
      background: var(--hw-card) !important;
    }}
    .card {{
      border-radius: var(--hw-radius) !important;
      border: 1px solid var(--hw-border) !important;
      box-shadow: var(--hw-shadow) !important;
      background: var(--hw-card) !important;
      overflow: hidden;
    }}
    .hw-surface-card {{ border-radius: var(--hw-radius) !important; }}
    .hw-hero {{
      background: var(--hw-card);
      border: 1px solid var(--hw-border);
      border-radius: var(--hw-radius);
      box-shadow: var(--hw-shadow);
      padding: 1.75rem 2rem;
      margin-bottom: 1.25rem;
      border-left: 4px solid var(--hw-accent);
    }}
    .hw-hero-title {{
      font-size: 2.35rem;
      font-weight: 800;
      color: var(--hw-text);
      margin: 0 0 0.35rem 0;
      letter-spacing: -0.03em;
      line-height: 1.15;
    }}
    .hw-hero-sub {{
      margin: 0;
      font-size: 0.95rem;
      color: var(--hw-muted);
      font-weight: 400;
    }}
    .hw-main-shell {{
      border-radius: var(--hw-radius);
      border: 1px solid var(--hw-border);
      background: var(--hw-card);
      box-shadow: var(--hw-shadow);
      overflow: hidden;
    }}
    .hw-panel-title {{
      padding: 1.1rem 1.35rem;
      border-bottom: 1px solid var(--hw-border);
      background: linear-gradient(180deg, #fafbfc 0%, #fff 100%);
    }}
    .hw-heading {{ font-size: 1.2rem; font-weight: 700; color: var(--hw-text); margin: 0; letter-spacing: -0.02em; }}
    .hw-subtitle {{ font-size: 0.88rem; color: var(--hw-muted); margin: 0.25rem 0 0 0; }}
    .nav-tabs {{ border-bottom: 1px solid var(--hw-border) !important; padding: 0 1rem; gap: 0.25rem; }}
    .nav-tabs .nav-link {{
      border: none !important;
      border-radius: var(--hw-radius) var(--hw-radius) 0 0 !important;
      color: var(--hw-muted) !important;
      font-weight: 500;
      padding: 0.65rem 1rem;
    }}
    .nav-tabs .nav-link.active {{
      color: var(--hw-accent) !important;
      background: transparent !important;
      border-bottom: 3px solid var(--hw-accent) !important;
      font-weight: 600;
    }}
    .nav-tabs .nav-link:hover {{ color: var(--hw-text) !important; }}
    .tab-content {{ padding: 1.35rem; background: var(--hw-card); }}
    .btn-primary {{
      background: var(--hw-accent) !important;
      border-color: var(--hw-accent) !important;
      border-radius: var(--hw-radius) !important;
      font-weight: 600;
    }}
    .btn-primary:hover {{ filter: brightness(0.95); }}
    .btn-outline-secondary {{
      border-radius: var(--hw-radius) !important;
      border-color: var(--hw-border) !important;
      color: var(--hw-text) !important;
    }}
    .alert {{
      border-radius: var(--hw-radius) !important;
      border: 1px solid var(--hw-border) !important;
    }}
    .hw-metric-card {{
      background: var(--hw-card);
      border: 1px solid var(--hw-border);
      border-radius: var(--hw-radius);
      padding: 1.35rem 1.5rem;
      box-shadow: var(--hw-shadow);
      min-height: 100%;
    }}
    .hw-metric-state-name {{ font-size: 1.35rem; font-weight: 700; color: var(--hw-text); letter-spacing: -0.02em; }}
    .hw-metric-hero {{ font-size: 2.6rem; font-weight: 800; color: var(--hw-accent); line-height: 1; letter-spacing: -0.03em; }}
    .hw-metric-hero-suffix {{ font-size: 1.25rem; font-weight: 700; color: var(--hw-accent); margin-left: 0.1rem; }}
    .hw-metric-caption {{ font-size: 0.8rem; color: var(--hw-muted); margin: 0; }}
    .hw-stat-label {{ font-size: 0.75rem; text-transform: uppercase; letter-spacing: 0.06em; color: var(--hw-muted); display: block; }}
    .hw-stat-value {{ font-size: 1.15rem; font-weight: 700; color: var(--hw-text); }}
    .hw-insight {{
      background: #eff6ff;
      border: 1px solid #bfdbfe;
      border-radius: var(--hw-radius);
      padding: 1rem 1.15rem;
      margin-bottom: 1.25rem;
    }}
    .hw-insight-title {{ font-size: 0.72rem; text-transform: uppercase; letter-spacing: 0.08em; font-weight: 700; color: var(--hw-accent); margin-bottom: 0.35rem; }}
    .hw-insight-body {{ font-size: 0.95rem; color: var(--hw-text); margin: 0; line-height: 1.5; }}
    .hw-section-label {{ font-size: 0.72rem; text-transform: uppercase; letter-spacing: 0.07em; font-weight: 600; color: var(--hw-muted); margin-bottom: 0.75rem; }}
    .hw-table-card .card-header {{
      background: var(--hw-card) !important;
      border-bottom: 1px solid var(--hw-border) !important;
      padding: 1rem 1.25rem !important;
    }}
    .hw-table-card h5, .hw-table-card h6 {{ color: var(--hw-text) !important; }}
    table.table {{ font-size: 0.9rem; }}
    .hw-page-wrap {{
      max-width: 1100px;
      margin-left: auto;
      margin-right: auto;
      padding-left: 1rem;
      padding-right: 1rem;
    }}
    .hw-input-card {{
      background: var(--hw-card);
      border: 1px solid var(--hw-border);
      border-radius: var(--hw-radius);
      box-shadow: var(--hw-shadow);
    }}
    .hw-results-placeholder {{
      min-height: 6rem;
    }}
    .hw-report-stats-wrap table tbody tr:nth-child(even) {{
      background: #f8f9fa;
    }}
    .hw-report-stats-wrap table tbody td {{
      padding: 0.5rem;
      border-bottom: 1px solid #eee;
      text-align: center;
    }}
    .hw-report-stats-wrap table tbody td:first-child {{
      text-align: left;
    }}
    """
    )


def _national_picture_tab_ui():
    """U.S. choropleth (data preloaded at startup)."""
    return ui.TagList(
        ui.h3(
            "U.S. mass shooting coverage by state (2025)",
            style=f"color: {_HW_TEXT}; font-weight: 700; margin-bottom: 0.5rem;",
        ),
        ui.p(
            "NYT coverage rate by state: GVA incidents with ≥1 keyword-matched NYT article in the cache window.",
            class_="hw-text-muted mb-4",
            style="font-size: 0.95rem; max-width: 42rem;",
        ),
        ui.output_ui("national_map"),
    )


def _compare_states_tab_ui():
    """Sidebar: State A/B, validation line, run button, cache, errors. Main: comparison results."""
    states = get_states_2025()
    choices = {s: s for s in states} if states else {"": "(No 2025 data)"}
    default_a = states[0] if states else None
    default_b = (states[1] if len(states) > 1 else states[0]) if states else None
    return ui.layout_sidebar(
        ui.sidebar(
            ui.p("Compare two states", class_="hw-section-label"),
            ui.div(
                ui.layout_columns(
                    ui.input_select("state_a", "State A", choices=choices, selected=default_a),
                    ui.input_select("state_b", "State B", choices=choices, selected=default_b),
                    col_widths=[6, 6],
                    style="gap: 1rem;",
                ),
                ui.div(
                    ui.output_text("validation_summary"),
                    class_="small mt-2 mb-0",
                    style=f"color: {_HW_MUTED}; line-height: 1.45;",
                ),
                ui.input_action_button(
                    "run_analysis",
                    "Run NYT coverage analysis",
                    class_="btn-primary mt-3",
                    width="100%",
                ),
                class_="hw-input-card p-3 p-md-4 mb-0",
            ),
            ui.output_ui("cache_status_ui"),
            ui.output_ui("analysis_error_ui"),
            width=340,
            open="desktop",
        ),
        ui.output_ui("comparison_section"),
        fillable=True,
    )


def _report_tab_ui():
    return ui.TagList(
        ui.p("AI pipeline report", class_="hw-section-label"),
        ui.p(
            "Four-agent pipeline (validation → demographics → analysis → narrative) for the "
            "two states selected in Compare States.",
            class_="hw-text-muted small mb-3",
            style="max-width: 48rem;",
        ),
        ui.row(
            ui.column(
                8,
                ui.input_action_button(
                    "go_to_report",
                    "Generate Report →",
                    class_="btn-primary mb-3",
                ),
                ui.output_ui("report_progress_ui"),
                ui.hr(),
                ui.output_ui("report_body"),
                ui.hr(),
                ui.output_ui("report_stats_table"),
            ),
            ui.column(
                4,
                ui.div(
                    ui.h5("About This Report", style="color: #1a3a5c;"),
                    ui.p(
                        "Generated by a 4-agent AI pipeline using "
                        "gpt-4o-mini. Demographics from U.S. Census "
                        "ACS 2023. Coverage data from NYT API + "
                        "Gun Violence Archive."
                    ),
                    ui.hr(),
                    ui.output_text("report_cost_info"),
                    ui.hr(),
                    ui.download_button(
                        "download_report_hw2",
                        "⬇ Download as Word",
                        class_="btn btn-outline-primary w-100",
                    ),
                    ui.hr(),
                    ui.h6("Data Sources"),
                    ui.tags.ul(
                        ui.tags.li("NYT Article Search API"),
                        ui.tags.li("Gun Violence Archive (GVA)"),
                        ui.tags.li("U.S. Census Bureau ACS 2023"),
                        ui.tags.li("OpenAI gpt-4o-mini"),
                    ),
                    style=(
                        "background: #f8f9fa; "
                        "padding: 1.25rem; "
                        "border-radius: 8px;"
                    ),
                ),
            ),
        ),
    )


app_ui = ui.page_fillable(
    ui.TagList(
        _app_global_css(),
        ui.div(
            ui.div(
                ui.h1("Who Gets Covered?", class_="hw-hero-title"),
                ui.p(
                    "Tracking New York Times coverage of mass shootings across America · 2025 · GVA + NYT API",
                    class_="hw-hero-sub",
                ),
                class_="hw-hero",
            ),
            ui.div(
                ui.output_ui("main_card_header"),
                ui.navset_tab(
                    ui.nav_panel("National Picture", _national_picture_tab_ui()),
                    ui.nav_panel("Compare States", _compare_states_tab_ui()),
                    ui.nav_panel("The Report", _report_tab_ui()),
                ),
                class_="hw-main-shell",
            ),
            class_="hw-page-wrap py-3",
        ),
    ),
    title="NYT Articles: State Comparison",
    fillable=True,
)

# 2. Server ############################


def server(input: Inputs, output: Outputs, session: Session):
    cache_articles = reactive.Value(None)
    cache_error = reactive.Value(None)
    analysis_result = reactive.Value(None)
    analysis_running = reactive.Value(False)
    report_rv = reactive.Value(None)  # Agent 4 text
    report_cost_rv = reactive.Value(None)  # cost dict
    report_progress_rv = reactive.Value({})  # agent step statuses
    stats_data_rv = reactive.Value(None)  # dict of pre-computed stats for Report tab UI
    outlier_data_rv = reactive.Value(None)  # list of outlier events for Report tab
    agent3_bullets_rv = reactive.Value(None)  # bullet strings only
    # Populated after NYT cache load so GVA↔NYT matching uses the same file as Compare States.
    hw2_events_rv = reactive.Value(None)
    national_stats_rv = reactive.Value(None)
    national_stats_err_rv = reactive.Value(None)

    _empty_national_cols = [
        "state_abbr",
        "state_name",
        "total_events",
        "covered_events",
        "coverage_rate",
        "outlier_count",
    ]

    @reactive.calc
    def national_map_bundle():
        ns = national_stats_rv.get()
        err = national_stats_err_rv.get()
        if ns is None and err is None:
            return _national_map_figure_bundle(pd.DataFrame(), is_loading=True)
        if err:
            return _national_map_figure_bundle(
                pd.DataFrame(columns=_empty_national_cols),
                load_error=err,
            )
        return _national_map_figure_bundle(ns.copy())

    @output
    @render.ui
    def national_map():
        return _plotly_figure_iframe_ui(national_map_bundle()["figure"])

    # Load cache once when app starts, then build national map + hw2_events from same JSON + GVA CSV.
    @reactive.effect
    def init_cache():
        if cache_articles.get() is not None or cache_error.get() is not None:
            return
        try:
            articles = load_or_build_2025_cache(
                cache_path=_HOMEWORK_1 / "nyt_2025_shootings_cache.json",
            )
            cache_articles.set(articles)
            cache_error.set(None)
            try:
                ev, _ = build_gva_events_for_states(_ALL_STATES_ABBR, max_articles=0)
                national_stats_rv.set(compute_national_stats(ev))
                hw2_events_rv.set(ev)
                national_stats_err_rv.set(None)
                print(
                    f"[HW2_app] National map: {len(ev)} GVA events; "
                    f"homework1_dir={homework1_dir()}",
                    flush=True,
                )
            except Exception as ex:
                national_stats_err_rv.set(str(ex))
                hw2_events_rv.set([])
                national_stats_rv.set(
                    pd.DataFrame(columns=_empty_national_cols),
                )
                print(f"[HW2_app] National stats / GVA match failed: {ex}", flush=True)
        except Exception as e:
            cache_articles.set(None)
            cache_error.set(str(e))
            national_stats_err_rv.set(f"NYT cache: {e}")
            hw2_events_rv.set([])
            national_stats_rv.set(pd.DataFrame(columns=_empty_national_cols))

    @output
    @render.ui
    def cache_status_ui():
        err = cache_error.get()
        if err:
            return ui.div(f"Cache error: {err}", class_="alert alert-danger mb-0")
        arts = cache_articles.get()
        if arts is None:
            return ui.div("Loading NYT articles cache…", class_="alert alert-info mb-0")
        return ui.div(f"NYT cache ready · {len(arts):,} articles loaded.", class_="alert alert-success mb-0")

    @output
    @render.text
    def validation_summary():
        try:
            state_a = input.state_a()
            state_b = input.state_b()
        except Exception:
            return "Select two states to begin."
        abbr_a = _state_pick_to_abbr(state_a)
        abbr_b = _state_pick_to_abbr(state_b)
        if not abbr_a or not abbr_b:
            return "Select two states to begin."

        ev_all = hw2_events_rv.get()
        if ev_all is None:
            return "Loading GVA / NYT event match data…"
        state_events = [
            e
            for e in ev_all
            if e.get("state") in [abbr_a, abbr_b]
        ]
        total = len(state_events)
        covered = sum(1 for e in state_events if e.get("matched_article_urls"))
        outliers = sum(1 for e in state_events if e.get("is_outlier"))

        lines = [
            f"{total} GVA events",
            f"{covered} matched NYT articles",
        ]
        if outliers > 0:
            lines.append(f"⚠️ {outliers} outlier event(s)")

        return " · ".join(lines)

    @output
    @render.ui
    def outlier_banner():
        try:
            state_a = input.state_a()
            state_b = input.state_b()
        except Exception:
            return ui.div()
        abbr_a = _state_pick_to_abbr(state_a)
        abbr_b = _state_pick_to_abbr(state_b)
        if not abbr_a or not abbr_b:
            return ui.div()

        ev_all = hw2_events_rv.get()
        if ev_all is None:
            return ui.div()
        outliers = [
            e
            for e in ev_all
            if e.get("state") in [abbr_a, abbr_b] and e.get("is_outlier")
        ]
        if not outliers:
            return ui.div()

        cities = ", ".join(
            f"{e.get('city', '?')} ({e.get('state', '?')})" for e in outliers
        )
        return ui.div(
            f"⚠️ High-profile outlier event(s) detected: {cities}. "
            f"Coverage rate for this state may be inflated. "
            f"The generated report will address this.",
            style=(
                "background: #FFF3CD; "
                "border: 1px solid #E8A838; "
                "border-radius: 4px; "
                "padding: 0.75rem 1rem; "
                "margin-bottom: 1rem; "
                "font-size: 0.9rem;"
            ),
        )

    @reactive.effect
    @reactive.event(input.go_to_report, ignore_none=False)
    def handle_generate_report():
        try:
            state_a = input.state_a()
            state_b = input.state_b()
        except Exception:
            return
        if not state_a or not state_b:
            return
        if state_a == state_b:
            report_progress_rv.set({})
            report_rv.set("Please select two different states in Compare States.")
            report_cost_rv.set(None)
            stats_data_rv.set(None)
            outlier_data_rv.set(None)
            agent3_bullets_rv.set(None)
            return

        abbr_a = _state_pick_to_abbr(state_a)
        abbr_b = _state_pick_to_abbr(state_b)
        if not abbr_a or not abbr_b:
            return

        report_progress_rv.set(
            {
                "agent1": "done",
                "agent2": "pending",
                "agent3": "pending",
                "agent4": "pending",
            }
        )
        report_rv.set(None)
        report_cost_rv.set(None)
        stats_data_rv.set(None)
        outlier_data_rv.set(None)
        agent3_bullets_rv.set(None)

        ev_all = hw2_events_rv.get()
        if ev_all is None:
            report_rv.set(
                "## Report unavailable\n\nGVA / NYT event data is still loading. "
                "Wait for the green NYT cache banner, then try again."
            )
            report_progress_rv.set(
                {
                    "agent1": "done",
                    "agent2": "done",
                    "agent3": "done",
                    "agent4": "done",
                }
            )
            return
        selected_events = [
            e
            for e in ev_all
            if e.get("state") in [abbr_a, abbr_b]
        ]
        validated_articles = _synthetic_validated_from_gva_events(selected_events)
        agent1_for_log = [
            {"url": v.get("url"), "relevant": v.get("relevant"), "reason": v.get("reason")}
            for v in validated_articles
        ]
        print("\n=== AGENT 1 OUTPUT (synthetic — pipeline input) ===", flush=True)
        print(json.dumps(agent1_for_log, indent=2), flush=True)

        report_progress_rv.set(
            {
                "agent1": "done",
                "agent2": "running",
                "agent3": "pending",
                "agent4": "pending",
            }
        )

        try:
            result = run_agent_pipeline(
                validated_articles=validated_articles,
                all_events=selected_events,
                state_a=abbr_a,
                state_b=abbr_b,
                rag_conn=rag_conn,
            )
        except Exception as e:
            report_progress_rv.set(
                {
                    "agent1": "done",
                    "agent2": "done",
                    "agent3": "done",
                    "agent4": "done",
                }
            )
            report_rv.set(f"## Report failed\n\n{str(e)}")
            report_cost_rv.set(None)
            stats_data_rv.set(None)
            outlier_data_rv.set(None)
            agent3_bullets_rv.set(None)
            return

        report_progress_rv.set(
            {
                "agent1": "done",
                "agent2": "done",
                "agent3": "done",
                "agent4": "done",
            }
        )

        report_rv.set(result.get("agent4", "") or "")
        report_cost_rv.set(result.get("tokens") or {})
        stats_data_rv.set(result.get("stats_data"))
        outlier_data_rv.set(result.get("outlier_events"))
        agent3_bullets_rv.set(result.get("agent3_bullets"))

    @output
    @render.ui
    def report_progress_ui():
        steps = report_progress_rv.get()
        if not steps:
            return ui.p(
                "Select two states in the Compare States tab, "
                "then click 'Generate Report →'.",
                style="color: #888; font-style: italic;",
            )
        icons = {"pending": "⏳", "running": "🔄", "done": "✅"}
        colors = {
            "pending": "#aaa",
            "running": "#1976D2",
            "done": "#2e7d32",
        }
        labels = {
            "agent1": "Agent 1: Article validation",
            "agent2": "Agent 2: Demographics enrichment",
            "agent3": "Agent 3: Pattern analysis",
            "agent4": "Agent 4: Report writing",
        }
        items = []
        for key, label in labels.items():
            status = steps.get(key, "pending")
            ic = icons.get(status, icons["pending"])
            items.append(
                ui.p(
                    f"{ic}  {label}",
                    style=(
                        f"color: {colors.get(status, colors['pending'])}; "
                        f"margin: 0.25rem 0; "
                        f"font-size: 0.95rem;"
                    ),
                )
            )
        return ui.div(
            *items,
            style=(
                "background: #f0f4f8; "
                "border-radius: 6px; "
                "padding: 1rem; "
                "margin-bottom: 1rem;"
            ),
        )

    @output
    @render.ui
    def report_body():
        text = report_rv.get()
        if not text:
            return ui.div()
        try:
            import markdown as md_lib

            html = md_lib.markdown(text)
        except ImportError:
            html = text.replace("\n", "<br>")
        return ui.div(
            ui.HTML(html),
            style="line-height: 1.7; font-size: 1rem;",
        )

    @output
    @render.ui
    def report_stats_table():
        stats = stats_data_rv.get()
        outliers = outlier_data_rv.get()
        bullets = agent3_bullets_rv.get()

        if not stats:
            return ui.div()

        sa = stats["state_a"]
        sb = stats["state_b"]
        sa_data = stats["state_a_data"]
        sb_data = stats["state_b_data"]
        demo_a = sa_data.get("demographics") if isinstance(sa_data.get("demographics"), dict) else {}
        demo_b = sb_data.get("demographics") if isinstance(sb_data.get("demographics"), dict) else {}

        def fmt_rate(v):
            if v is None:
                return "N/A"
            try:
                return f"{float(v):.1f}%"
            except (TypeError, ValueError):
                return "N/A"

        def fmt_currency(v):
            if v is None:
                return "N/A"
            try:
                return f"${float(v):,.0f}"
            except (TypeError, ValueError):
                return "N/A"

        def fmt_days(v):
            if v is None:
                return "N/A"
            try:
                return f"{float(v):.1f} days"
            except (TypeError, ValueError):
                return "N/A"

        def fmt_pop(v):
            if v is None:
                return "N/A"
            try:
                return f"{float(v):,.0f}"
            except (TypeError, ValueError):
                return "N/A"

        cov_a = float(sa_data.get("coverage_rate") or 0)
        cov_b = float(sb_data.get("coverage_rate") or 0)
        sa_higher = cov_a > cov_b
        sb_higher = cov_b > cov_a

        rows = [
            ("Total Shootings", str(sa_data.get("total_events", "")), str(sb_data.get("total_events", ""))),
            ("NYT Covered", str(sa_data.get("covered_events", "")), str(sb_data.get("covered_events", ""))),
            ("Coverage Rate", fmt_rate(sa_data.get("coverage_rate")), fmt_rate(sb_data.get("coverage_rate"))),
            (
                "Avg Days to First Article",
                fmt_days(sa_data.get("avg_days_to_first_article")),
                fmt_days(sb_data.get("avg_days_to_first_article")),
            ),
            (
                "Same-Day Coverage",
                fmt_rate(sa_data.get("same_day_coverage_pct")),
                fmt_rate(sb_data.get("same_day_coverage_pct")),
            ),
            (
                "Median Household Income",
                fmt_currency(demo_a.get("median_household_income")),
                fmt_currency(demo_b.get("median_household_income")),
            ),
            (
                "% White Population",
                fmt_rate(demo_a.get("pct_white")),
                fmt_rate(demo_b.get("pct_white")),
            ),
            (
                "State Population",
                fmt_pop(demo_a.get("population")),
                fmt_pop(demo_b.get("population")),
            ),
        ]

        table_rows = []
        for metric, val_a, val_b in rows:
            if metric == "Coverage Rate":
                cell_a = ui.tags.td(
                    ui.tags.b(val_a) if sa_higher else val_a,
                    style=("color: #2e7d32;" if sa_higher else "color: #888;"),
                )
                cell_b = ui.tags.td(
                    ui.tags.b(val_b) if sb_higher else val_b,
                    style=("color: #2e7d32;" if sb_higher else "color: #888;"),
                )
            else:
                cell_a = ui.tags.td(val_a)
                cell_b = ui.tags.td(val_b)
            table_rows.append(
                ui.tags.tr(
                    ui.tags.td(metric, style="color: #555; font-size: 0.9rem;"),
                    cell_a,
                    cell_b,
                )
            )

        corr = stats.get("correlation_flag") or ""
        cl = corr.lower()
        if "negative" in cl:
            badge_color = "#E8A838"
            badge_text = "⚠️ Negative correlation detected"
            badge_subtext = "Higher income state has LOWER coverage rate"
        elif "positive" in cl:
            badge_color = "#2196F3"
            badge_text = "📈 Positive correlation"
            badge_subtext = "Higher income state also has higher coverage rate"
        else:
            badge_color = "#888"
            badge_text = "Coverage comparison"
            badge_subtext = ""

        outlier_section = ui.div()
        if outliers:
            outlier_items = [
                ui.tags.li(
                    f"{e.get('city', '?')}, {e.get('state', '?')} "
                    f"({e.get('date', '?')}) — "
                    f"{e.get('article_count', 0)} NYT articles matched",
                    style="margin-bottom: 0.25rem;",
                )
                for e in outliers
            ]
            outlier_section = ui.div(
                ui.h6(
                    "⚠️ High-Profile Outlier Events Detected",
                    style="color: #E8A838; margin-bottom: 0.5rem;",
                ),
                ui.p(
                    "These events received 5+ matched articles and may inflate the state's coverage rate. "
                    "The report accounts for this.",
                    style="font-size: 0.85rem; color: #666;",
                ),
                ui.tags.ul(*outlier_items, style="font-size: 0.9rem;"),
                style=(
                    "background: #FFF8E1; "
                    "border: 1px solid #E8A838; "
                    "border-radius: 6px; "
                    "padding: 1rem; "
                    "margin-top: 1rem;"
                ),
            )

        bullets_section = ui.div()
        if bullets:
            bullet_items = [ui.tags.li(b, style="margin-bottom: 0.5rem;") for b in bullets]
            bullets_section = ui.div(
                ui.h6("Key Findings", style="color: #1a3a5c; margin-bottom: 0.75rem;"),
                ui.tags.ul(
                    *bullet_items,
                    style="font-size: 0.95rem; line-height: 1.6;",
                ),
                style="margin-top: 1.5rem;",
            )

        speed = stats.get("speed_note") or ""
        speed_block = (
            ui.p(
                f"⏱ {speed}",
                style="font-size: 0.9rem; color: #555; margin-top: 0.75rem;",
            )
            if speed
            else ui.div()
        )

        return ui.div(
            ui.h5("Statistical Summary", style="color: #1a3a5c; margin-bottom: 1rem;"),
            ui.div(
                ui.tags.table(
                    ui.tags.thead(
                        ui.tags.tr(
                            ui.tags.th(
                                "Metric",
                                style=(
                                    "text-align: left; padding: 0.5rem; "
                                    "border-bottom: 2px solid #1a3a5c; width: 45%;"
                                ),
                            ),
                            ui.tags.th(
                                sa,
                                style=(
                                    "text-align: center; padding: 0.5rem; "
                                    "border-bottom: 2px solid #1a3a5c; color: #1a3a5c;"
                                ),
                            ),
                            ui.tags.th(
                                sb,
                                style=(
                                    "text-align: center; padding: 0.5rem; "
                                    "border-bottom: 2px solid #1a3a5c; color: #1a3a5c;"
                                ),
                            ),
                        )
                    ),
                    ui.tags.tbody(*table_rows),
                    style="width: 100%; border-collapse: collapse; font-size: 0.95rem;",
                ),
                style=(
                    "background: white; border-radius: 6px; padding: 1rem; "
                    "box-shadow: 0 1px 3px rgba(0,0,0,0.06);"
                ),
            ),
            ui.div(
                ui.p(badge_text, style="font-weight: 600; margin: 0;"),
                ui.p(badge_subtext, style="font-size: 0.85rem; color: #555; margin: 0;")
                if badge_subtext
                else ui.div(),
                style=(
                    f"background: {badge_color}22; border-left: 4px solid {badge_color}; "
                    "padding: 0.75rem 1rem; border-radius: 4px; margin-top: 1rem;"
                ),
            ),
            speed_block,
            outlier_section,
            bullets_section,
            class_="hw-report-stats-wrap",
            style="margin-top: 1.5rem;",
        )

    @output
    @render.text
    def report_cost_info():
        tokens = report_cost_rv.get()
        if not tokens:
            return ""
        total = int(tokens.get("prompt", 0) or 0) + int(tokens.get("completion", 0) or 0)
        cost = total / 1_000_000 * 0.60
        return f"{total:,} tokens · ${cost:.4f}"

    def _hw2_report_docx_filename():
        try:
            sa = (input.state_a() or "State").replace(" ", "_")
            sb = (input.state_b() or "State").replace(" ", "_")
        except Exception:
            sa, sb = "State", "State"
        return f"shooting_coverage_report_{sa}_{sb}.docx"

    @output
    @render.download(filename=_hw2_report_docx_filename)
    def download_report_hw2():
        import os
        import tempfile

        text = report_rv.get()
        if not text:
            raise RuntimeError("Generate the report on The Report tab first.")
        try:
            sa = input.state_a() or "State A"
            sb = input.state_b() or "State B"
        except Exception:
            sa, sb = "State A", "State B"
        title = f"Shooting coverage report: {sa} vs {sb}"
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            _write_hw2_pipeline_docx(
                Path(path),
                title,
                text,
                stats_data=stats_data_rv.get(),
                bullet_points=agent3_bullets_rv.get(),
                outlier_events=outlier_data_rv.get(),
            )
            return path
        except Exception:
            try:
                os.unlink(path)
            except OSError:
                pass
            raise

    @output
    @render.ui
    def main_card_header():
        return ui.div(
            ui.h4("State comparison", class_="hw-heading"),
            ui.p("Coverage rates and incident counts for your two selected states.", class_="hw-subtitle"),
            class_="hw-panel-title",
        )

    @reactive.effect
    @reactive.event(input.run_analysis, ignore_none=False)
    def run_analysis():
        analysis_running.set(True)
        analysis_result.set(None)
        try:
            err = cache_error.get()
            if err:
                analysis_result.set({"error": err, "state_a": None, "state_b": None})
                return
            arts = cache_articles.get()
            if arts is None:
                analysis_result.set({"error": "Cache not ready yet.", "state_a": None, "state_b": None})
                return
            try:
                sa = input.state_a()
                sb = input.state_b()
            except Exception:
                analysis_result.set({"error": "Select both states.", "state_a": None, "state_b": None})
                return
            if not sa or not sb:
                analysis_result.set({"error": "Select both states.", "state_a": None, "state_b": None})
                return
            if sa == sb:
                analysis_result.set({"error": "Please select two different states.", "state_a": None, "state_b": None})
                return
            with Progress(min=0, max=2, session=session) as p:
                p.set(0, message="Loading articles…", detail="Using cached NYT corpus")
                p.set(1, message="Computing statistics…", detail="Matching GVA incidents to NYT coverage")
                result = run_state_analysis(sa, sb, cache_articles=arts)
                analysis_result.set(result)
                p.set(2, message="Complete", detail="Results are ready below.")
        finally:
            analysis_running.set(False)

    @output
    @render.ui
    def analysis_error_ui():
        r = analysis_result.get()
        if r is None:
            return None
        if r.get("error"):
            return ui.div(r["error"], class_="alert alert-warning mb-0")
        return None

    @output
    @render.ui
    def comparison_section():
        if analysis_running.get():
            return ui.div(
                ui.p("Results", class_="hw-section-label"),
                ui.div(
                    ui.p(
                        "Running analysis…",
                        class_="hw-text-muted small text-center mb-0",
                    ),
                    class_="hw-results-placeholder d-flex align-items-center justify-content-center py-5",
                    style=(
                        f"border: 1px dashed {_HW_BORDER}; border-radius: {_HW_RADIUS}; "
                        "background: #fafbfc;"
                    ),
                ),
                class_="mb-2",
            )
        r = analysis_result.get()
        if r is None:
            return ui.div(
                ui.p("Results", class_="hw-section-label"),
                ui.h5("Run an analysis", class_="mb-2", style=f"color: {_HW_TEXT}; font-weight: 700; font-size: 1.25rem;"),
                ui.p(
                    "Choose two states above, then click “Run NYT coverage analysis”.",
                    class_="hw-text-muted mb-0",
                    style="max-width: 36rem; line-height: 1.5;",
                ),
                class_="py-1",
            )
        if r.get("error"):
            return None
        a = r.get("state_a") or {}
        b = r.get("state_b") or {}
        sa_name = a.get("state", "State A")
        sb_name = b.get("state", "State B")
        cache_start = r.get("cache_start_date", "")
        cache_end = r.get("cache_end_date", "")
        pa = float(a.get("pct_reported") or 0)
        pb = float(b.get("pct_reported") or 0)
        if pa > pb:
            insight = (
                f"{sa_name} has a higher share of incidents with NYT coverage ({pa:.0f}% vs {pb:.0f}%). "
                f"Interpret with care: one newspaper cannot represent all media attention."
            )
        elif pb > pa:
            insight = (
                f"{sb_name} has a higher share of incidents with NYT coverage ({pb:.0f}% vs {pa:.0f}%). "
                f"Interpret with care: one newspaper cannot represent all media attention."
            )
        else:
            insight = (
                f"Both states show the same NYT coverage rate ({pa:.0f}%). "
                f"See the table below for incident counts and coverage timing."
            )
        insight_block = ui.div(
            ui.p("At a glance", class_="hw-insight-title"),
            ui.p(insight, class_="hw-insight-body"),
            class_="hw-insight mb-4",
        )

        head = [
            ui.p("Comparison results", class_="hw-section-label"),
        ]
        if cache_start and cache_end:
            head.append(
                ui.p(
                    ui.span("GVA window: ", class_="fw-semibold", style=f"color: {_HW_TEXT};"),
                    ui.span(f"{cache_start} – {cache_end}", class_="hw-text-muted"),
                    ui.span(" · NYT cache article dates", class_="small hw-text-muted ms-1"),
                    class_="small mb-3",
                )
            )
        head.append(insight_block)

        return ui.TagList(
            *head,
            ui.layout_columns(
                _state_summary_card(sa_name, a),
                _state_summary_card(sb_name, b),
                col_widths=[6, 6],
                style="gap: 1rem;",
            ),
            ui.card(
                ui.card_header(ui.h5("Full metrics", class_="mb-0", style=f"font-size: 1rem; font-weight: 700; color: {_HW_TEXT};")),
                ui.card_body(
                    ui.output_ui("outlier_banner"),
                    ui.output_table("comparison_table"),
                    class_="p-3 pt-0",
                ),
                class_="mt-4 hw-table-card",
            ),
            ui.div(
                ui.p("Coverage detail by event", class_="hw-section-label mt-4"),
                ui.layout_columns(
                    ui.div(_takeaway_card(sa_name, a), class_="mb-0"),
                    ui.div(_takeaway_card(sb_name, b), class_="mb-0"),
                    col_widths=[6, 6],
                    style="gap: 1rem;",
                ),
            ),
        )

    @output
    @render.table
    def comparison_table():
        r = analysis_result.get()
        if r is None or r.get("error"):
            return None
        rows = _comparison_table_rows(r)
        a = r.get("state_a") or {}
        b = r.get("state_b") or {}
        import pandas as pd
        df = pd.DataFrame(rows, columns=["Metric", a.get("state", "State A"), b.get("state", "State B")])
        return df

# 3. App ############################

app = App(app_ui, server)
