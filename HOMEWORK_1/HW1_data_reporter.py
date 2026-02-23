# HW1_data_reporter.py
# Two-state NYT 2025 coverage report: AI comparative analysis and comparison table (.docx).
# Uses HW1_nyt_cache and HW1_state_analysis. Optional OpenAI for deeper analytical sections.

import json
import os
from pathlib import Path
from typing import Optional

from docx import Document

from HW1_nyt_cache import load_or_build_2025_cache
from HW1_state_analysis import run_state_analysis

HW1_DIR = Path(__file__).resolve().parent
REPO_ROOT = HW1_DIR.parent
FONT_NAME = "Calibri"

# Optional: load OpenAI API key for AI-generated analysis
try:
    from dotenv import load_dotenv
    load_dotenv(REPO_ROOT / ".env")
except ImportError:
    pass
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")

AI_COMPARISON_SYSTEM = """You are an academic researcher writing a comparative analysis of New York Times coverage of mass shootings in two US states. Apply these rules strictly:
- Do not speculate about causes. Do not introduce external information. Base all conclusions strictly on the data provided.
- Write in a formal academic tone. Avoid repetition. Interpret patterns rather than restating values.
- Cite specific numeric values when making comparisons. Use whole numbers only (no decimal points).
- Do not make political or normative claims."""

AI_COMPARISON_USER_TEMPLATE = """Using only the following comparison data for {state_a_name} and {state_b_name}, produce an analytical report. Do not speculate about causes or introduce external information.

**Queried data:** Shootings and NYT articles are restricted to the date range {date_range}.

**Data (from GVA and NYT article cache; whole numbers only):**
{data_json}

**Output requirements:**
Respond in valid JSON with exactly four keys (no other keys, no text outside the JSON). Use whole numbers only when citing statistics.

1. "executive_summary": string — Brief executive summary (one or two short paragraphs). Interpretive, not descriptive. Synthesize main comparative findings without restating the table. Do not repeat this content in comparative_findings.

2. "comparative_findings": string — Non-repetitive comparative findings in three parts. Use clear subheadings or paragraph breaks for (a), (b), (c). Key observations only; interpret patterns.
   (a) Reporting Likelihood: Compare total events and percent reported between the two states. Calculate and interpret absolute and relative differences (e.g., difference in percentage points; ratio of percentages where meaningful). Comment on denominator size effects (how total event count affects interpretability).
   (b) Coverage Intensity: Compare duration statistics (min, max, mean, median). Interpret whether coverage was brief or sustained in each state and what the spread indicates.
   (c) Distribution Pattern: Assess whether coverage is concentrated in a single event or spread across many. Explain how this affects summary statistics (e.g., mean vs median, stability of estimates).

3. "methodological_considerations": string — Very brief: maximum 5 sentences. Discuss small sample sizes where relevant, keyword-based matching limitations, and avoid causal inference. Formal tone.

4. "recommendation": string — One short paragraph: how the reader can get more information on or support efforts related to gun violence in America. Neutral, factual tone (e.g., data sources, evidence-based resources). No political or normative claims.

Do not include markdown code fences around the JSON. Output only the JSON object."""


def _set_paragraph_font(paragraph, font_name=FONT_NAME):
    """Set font for all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name


def _fmt(val):
    """Format value for table; no decimals. Use em dash if None."""
    if val is None:
        return "—"
    if isinstance(val, float):
        return str(int(round(val)))
    return str(val)


def _round_stat(v):
    """Round to integer for display; None stays None."""
    if v is None:
        return None
    try:
        return int(round(float(v)))
    except (TypeError, ValueError):
        return v


def _comparison_data_dict(
    state_a_name: str,
    state_b_name: str,
    state_a: dict,
    state_b: dict,
    cache_start_date: Optional[str] = None,
    cache_end_date: Optional[str] = None,
) -> dict:
    """Build a serializable dict of comparison stats for the AI prompt (integers, no decimals)."""
    out = {
        "queried_data_date_range": {
            "start_date": cache_start_date,
            "end_date": cache_end_date,
        },
        state_a_name: {
            "total_shootings": _round_stat(state_a.get("total_shootings")),
            "events_with_nyt_article": _round_stat(state_a.get("reported_count")),
            "pct_reported": _round_stat(state_a.get("pct_reported")),
            "coverage_duration_days": {
                "min": _round_stat(state_a.get("coverage_min_days")),
                "max": _round_stat(state_a.get("coverage_max_days")),
                "mean": _round_stat(state_a.get("coverage_mean_days")),
                "median": _round_stat(state_a.get("coverage_median_days")),
            },
        },
        state_b_name: {
            "total_shootings": _round_stat(state_b.get("total_shootings")),
            "events_with_nyt_article": _round_stat(state_b.get("reported_count")),
            "pct_reported": _round_stat(state_b.get("pct_reported")),
            "coverage_duration_days": {
                "min": _round_stat(state_b.get("coverage_min_days")),
                "max": _round_stat(state_b.get("coverage_max_days")),
                "mean": _round_stat(state_b.get("coverage_mean_days")),
                "median": _round_stat(state_b.get("coverage_median_days")),
            },
        },
    }
    return out


def _get_ai_comparison_analysis(
    state_a_name: str,
    state_b_name: str,
    state_a: dict,
    state_b: dict,
    cache_start_date: Optional[str] = None,
    cache_end_date: Optional[str] = None,
) -> Optional[dict]:
    """
    Call OpenAI to generate Executive Summary (with comparative analysis), Methodological Limitations, and Recommendation.
    Returns dict with keys executive_summary, methodological_limitations, recommendation, or None on failure.
    """
    if not OPENAI_API_KEY:
        return None
    try:
        import requests
    except ImportError:
        return None
    date_range = f"{cache_start_date or 'N/A'} to {cache_end_date or 'N/A'}"
    data = _comparison_data_dict(
        state_a_name, state_b_name, state_a, state_b,
        cache_start_date=cache_start_date,
        cache_end_date=cache_end_date,
    )
    data_json = json.dumps(data, indent=2)
    user_prompt = AI_COMPARISON_USER_TEMPLATE.format(
        state_a_name=state_a_name,
        state_b_name=state_b_name,
        date_range=date_range,
        data_json=data_json,
    )
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json",
    }
    body = {
        "model": "gpt-4o-mini",
        "messages": [
            {"role": "system", "content": AI_COMPARISON_SYSTEM},
            {"role": "user", "content": user_prompt},
        ],
        "response_format": {"type": "json_object"},
        "temperature": 0.3,
    }
    try:
        response = requests.post(url, headers=headers, json=body, timeout=90)
        response.raise_for_status()
        result = response.json()
        text = (result.get("choices") or [{}])[0].get("message", {}).get("content", "")
        if not text:
            return None
        text = text.strip()
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
        parsed = json.loads(text)
        return {
            "executive_summary": parsed.get("executive_summary", ""),
            "comparative_findings": parsed.get("comparative_findings", ""),
            "methodological_considerations": parsed.get("methodological_considerations", ""),
            "recommendation": parsed.get("recommendation", ""),
        }
    except (requests.exceptions.RequestException, json.JSONDecodeError, KeyError, IndexError):
        return None


def _executive_summary_text(state_a: dict, state_b: dict) -> str:
    """Fallback short narrative summary when AI is not used."""
    sa = state_a.get("state", "State A")
    sb = state_b.get("state", "State B")
    ta = state_a.get("total_shootings", 0)
    tb = state_b.get("total_shootings", 0)
    ra = state_a.get("reported_count", 0)
    rb = state_b.get("reported_count", 0)
    pa = int(round(state_a.get("pct_reported") or 0))
    pb = int(round(state_b.get("pct_reported") or 0))
    text = (
        f"This report compares New York Times coverage of mass shootings in 2025 for {sa} and {sb}, "
        "using Gun Violence Archive (GVA) events and a cached set of NYT articles from 2025. "
        f"{sa} had {ta} mass shooting events in 2025, of which {ra} ({pa}%) had at least one matching NYT article. "
        f"{sb} had {tb} events, with {rb} ({pb}%) reported by the NYT. "
    )
    mean_a = state_a.get("coverage_mean_days")
    mean_b = state_b.get("coverage_mean_days")
    if mean_a is not None or mean_b is not None:
        text += (
            "Among reported events, coverage duration (days from first to last article) varied by state; "
            "see the statistical comparison table below for min, max, mean, and median. "
        )
    text += (
        "Matching was based on city name and shooting-related keywords in headline or abstract, "
        "with articles restricted to publication in 2025."
    )
    return text


def generate_comparison_docx(
    state_a_name: str,
    state_b_name: str,
    state_a: dict,
    state_b: dict,
    output_path: Path,
    cache_start_date: Optional[str] = None,
    cache_end_date: Optional[str] = None,
) -> None:
    """
    Write a .docx with title, AI-generated Executive Summary (with comparative analysis),
    brief Methodological Limitations, Recommendation, and Statistical Comparison table (whole numbers).
    Falls back to non-AI summary if OPENAI_API_KEY is not set or the API call fails.
    """
    doc = Document()
    title = f"NYT 2025 Mass Shooting Coverage: {state_a_name} vs {state_b_name}"
    h0 = doc.add_heading(title, level=0)
    _set_paragraph_font(h0)

    ai_result = _get_ai_comparison_analysis(
        state_a_name, state_b_name, state_a, state_b,
        cache_start_date=cache_start_date,
        cache_end_date=cache_end_date,
    )

    # Executive Summary (interpretive, includes comparative analysis when AI is used)
    h1 = doc.add_heading("Executive Summary", level=1)
    _set_paragraph_font(h1)
    if ai_result and ai_result.get("executive_summary"):
        for para in ai_result["executive_summary"].strip().split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                _set_paragraph_font(p)
    else:
        p1 = doc.add_paragraph(_executive_summary_text(state_a, state_b))
        _set_paragraph_font(p1)

    # Comparative Findings (reporting likelihood, coverage intensity, distribution pattern)
    h2 = doc.add_heading("Comparative Findings", level=1)
    _set_paragraph_font(h2)
    if ai_result and ai_result.get("comparative_findings"):
        for block in ai_result["comparative_findings"].strip().split("\n\n"):
            if block.strip():
                p = doc.add_paragraph(block.strip())
                _set_paragraph_font(p)
    else:
        p_cf = doc.add_paragraph(
            "Comparative findings are based on the statistical comparison table below. "
            "Consider reporting likelihood (total events, percent reported, absolute and relative differences, denominator effects), "
            "coverage intensity (duration statistics and whether coverage was brief or sustained), "
            "and distribution pattern (concentration in few events vs spread). Set OPENAI_API_KEY for AI-generated analysis."
        )
        _set_paragraph_font(p_cf)

    # Methodological Considerations (very brief, max 5 sentences)
    h3 = doc.add_heading("Methodological Considerations", level=1)
    _set_paragraph_font(h3)
    if ai_result and ai_result.get("methodological_considerations"):
        p_mc = doc.add_paragraph(ai_result["methodological_considerations"].strip())
        _set_paragraph_font(p_mc)
    else:
        p_mc = doc.add_paragraph(
            "Small sample sizes in some states limit precision. Matching relies on keyword and city in headline or abstract; "
            "each article is allocated to one event. No causal inference is implied."
        )
        _set_paragraph_font(p_mc)

    # Recommendation: more information / support
    h4 = doc.add_heading("Further Information and Support", level=1)
    _set_paragraph_font(h4)
    if ai_result and ai_result.get("recommendation"):
        p_rec = doc.add_paragraph(ai_result["recommendation"].strip())
        _set_paragraph_font(p_rec)
    else:
        p_rec = doc.add_paragraph(
            "For more data on gun violence in the United States, see the Gun Violence Archive (gva.org). "
            "Readers seeking information on advocacy or support can search for evidence-based resources and organizations."
        )
        _set_paragraph_font(p_rec)

    # Statistical Comparison table (no decimal points)
    h5 = doc.add_heading("Statistical Comparison", level=1)
    _set_paragraph_font(h5)
    rows = [
        ("Total mass shootings (2025)", _fmt(state_a.get("total_shootings")), _fmt(state_b.get("total_shootings"))),
        ("Events with ≥1 NYT article", _fmt(state_a.get("reported_count")), _fmt(state_b.get("reported_count"))),
        ("% reported by NYT", f"{_fmt(state_a.get('pct_reported'))}%", f"{_fmt(state_b.get('pct_reported'))}%"),
        ("Coverage duration (days) — min", _fmt(state_a.get("coverage_min_days")), _fmt(state_b.get("coverage_min_days"))),
        ("Coverage duration (days) — max", _fmt(state_a.get("coverage_max_days")), _fmt(state_b.get("coverage_max_days"))),
        ("Coverage duration (days) — mean", _fmt(state_a.get("coverage_mean_days")), _fmt(state_b.get("coverage_mean_days"))),
        ("Coverage duration (days) — median", _fmt(state_a.get("coverage_median_days")), _fmt(state_b.get("coverage_median_days"))),
    ]
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = state_a_name
    table.rows[0].cells[2].text = state_b_name
    for i, (metric, va, vb) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = va
        table.rows[i].cells[2].text = vb
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _set_paragraph_font(paragraph)
    doc.save(output_path)


def run_report(state_a: str, state_b: str, output_dir: Path = None) -> tuple:
    """
    Full pipeline: load cache, run state analysis, write comparison .docx.
    Returns (success: bool, message: str). message is output path on success or error string.
    """
    if not state_a or not state_b:
        return False, "Please provide both state names."
    if state_a == state_b:
        return False, "Please provide two different states."
    out_dir = Path(output_dir) if output_dir else HW1_DIR
    safe = f"{state_a.replace(' ', '_')}_vs_{state_b.replace(' ', '_')}"
    output_path = out_dir / f"HW1_NYT_Comparison_{safe}.docx"

    try:
        cache = load_or_build_2025_cache()
    except (ValueError, RuntimeError) as e:
        return False, f"Cache error: {e}"

    result = run_state_analysis(state_a, state_b, cache_articles=cache)
    if result.get("error"):
        return False, f"State analysis error: {result['error']}"
    a = result.get("state_a")
    b = result.get("state_b")
    if not a or not b:
        return False, "State analysis returned no data."

    try:
        generate_comparison_docx(
            state_a_name=state_a,
            state_b_name=state_b,
            state_a=a,
            state_b=b,
            output_path=output_path,
            cache_start_date=result.get("cache_start_date"),
            cache_end_date=result.get("cache_end_date"),
        )
    except OSError as e:
        return False, f"Could not write report: {e}"
    except Exception as e:
        return False, f"Report generation failed: {e}"

    return True, str(output_path)


if __name__ == "__main__":
    import sys
    states_2025 = None
    try:
        from HW1_state_analysis import get_states_2025
        states_2025 = get_states_2025()
    except Exception:
        pass
    if len(sys.argv) >= 3:
        state_a = sys.argv[1]
        state_b = sys.argv[2]
    elif states_2025 and len(states_2025) >= 2:
        state_a, state_b = states_2025[0], states_2025[1]
        print(f"Using default states: {state_a}, {state_b}")
    else:
        print("Usage: python HW1_data_reporter.py <StateA> <StateB>")
        sys.exit(1)
    print("Running HW1 Data Reporter...")
    ok, msg = run_report(state_a, state_b)
    if ok:
        print(f"Report saved: {msg}")
    else:
        print(f"Error: {msg}")
        sys.exit(1)
